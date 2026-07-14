# rag/parsers.py
"""
ШАГ 3 в pipeline ingestion.

Парсеры извлекают чистый текст из разных форматов файлов.

СТРАТЕГИЯ: вытаскиваем только текст, всё остальное выбрасываем.
Цвета, шрифты, отступы, изображения — в векторном поиске бесполезны.

ПОДДЕРЖИВАЕМЫЕ ФОРМАТЫ:
- .txt  → читаем как есть (UTF-8)
- .pdf  → pymupdf (fitz): быстрый, хорошо держит кириллицу
           ВАЖНО: работает только с текстовыми PDF. Сканы (картинки) вернут пустоту.
- .docx → python-docx: параграфы + таблицы (таблицы важны — в них часто характеристики товаров)
- .xlsx → openpyxl: все листы, ячейки построчно через " | "
"""
from __future__ import annotations
from pathlib import Path


def parse_file(path: Path) -> str:
    """
    Диспетчер: определяет формат по расширению, вызывает нужный парсер.
    Возвращает сырой текст для дальнейшего чанкинга.
    """
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return _parse_txt(path)
    elif suffix == ".pdf":
        return _parse_pdf(path)
    elif suffix == ".docx":
        return _parse_docx(path)
    elif suffix in (".xlsx", ".xls"):
        return _parse_xlsx(path)
    else:
        raise ValueError(f"Неподдерживаемый формат: {suffix}")


def _parse_txt(path: Path) -> str:
    # errors="replace" — не падаем на кривой кодировке, заменяем символ на "?"
    return path.read_text(encoding="utf-8", errors="replace")


def _parse_pdf(path: Path) -> str:
    import fitz  # pip install pymupdf

    doc = fitz.open(str(path))
    pages: list[str] = []
    for page_num, page in enumerate(doc):
        # "text" — plain text без разметки. Альтернатива "html" даёт больше структуры,
        # но нам нужен чистый текст для embedding.
        text = page.get_text("text")
        if text.strip():
            pages.append(text)
    doc.close()

    # Разделяем страницы двойным переносом — chunker увидит их как отдельные абзацы
    return "\n\n".join(pages)


def _parse_docx(path: Path) -> str:
    from docx import Document  # pip install python-docx

    doc = Document(str(path))
    parts: list[str] = []

    # Параграфы — основной текст документа
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Таблицы — часто содержат технические характеристики товаров
    # Формат: "Колонка1 | Колонка2 | Колонка3" для каждой строки
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def _parse_xlsx(path: Path) -> str:
    from openpyxl import load_workbook  # pip install openpyxl

    # read_only=True — не грузим всё в память сразу
    # data_only=True — берём значения ячеек, а не формулы (нам нужны данные)
    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []

    for sheet in wb.worksheets:
        sheet_parts: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            # values_only=True — возвращает значения ячеек, не Cell-объекты
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                sheet_parts.append(" | ".join(cells))

        if sheet_parts:
            # Добавляем название листа как контекст
            parts.append(f"[{sheet.title}]\n" + "\n".join(sheet_parts))

    wb.close()
    return "\n\n".join(parts)